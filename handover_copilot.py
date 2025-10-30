"""
Simple Handover Q&A – single-file MVP (stable, quote-safe, Excel support)
----------------------------------------------------------------------
• Ingest: PDF / TXT / MD / DOCX / DRAWIO / (optional OCR: PNG/JPG) / XLSX / XLS
• Stores FULL chunk text -> better offline answers (sentences + citations)
• Embeddings: OpenAI (retry) or local SentenceTransformers (auto fallback)
• Q&A: OpenAI chat when available, else offline extractive
• Rebuild flag: `ingest <folder> --rebuild`
• Env:
  - HANDOVER_FORCE_LOCAL=1  → never call OpenAI (embeddings or chat)
  - HF_* tokens ignored for public models to avoid 401
"""
from __future__ import annotations

import os, re, sys, json, time, shutil, hashlib, pathlib, argparse, logging
from typing import List, Dict, Any, Tuple, Optional

logging.basicConfig(level=logging.INFO, format="[%(levelname)s] %(message)s")

# Paths
STORE_DIR = pathlib.Path("./store")
NODES_PATH = STORE_DIR / "nodes.jsonl"
EMB_PATH   = STORE_DIR / "embeddings.npy"

# Deps
from dotenv import load_dotenv

try:
    from openai import OpenAI
except Exception:
    OpenAI = None  # guarded usage

from pypdf import PdfReader

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False

try:
    import docx  # python-docx
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

try:
    from sentence_transformers import SentenceTransformer
    LOCAL_EMBED_AVAILABLE = True
except Exception:
    LOCAL_EMBED_AVAILABLE = False

import numpy as np
from tqdm import tqdm
import xml.etree.ElementTree as ET

SUPPORTED_EXTS = {
    ".pdf", ".txt", ".md", ".png", ".jpg", ".jpeg", ".docx", ".drawio",
    ".xlsx", ".xls"
}
IMG_EXTS = {".png", ".jpg", ".jpeg"}

# Globals
client: Optional[OpenAI] = None
LOCAL_EMBEDDINGS = False
local_model = None
EMBED_DIM = 1536  # OpenAI default; switches to 384 for local

# ---------- Utils ----------
def sha1(s: str) -> str:
    h = hashlib.sha1()
    h.update(s.encode("utf-8", errors="ignore"))
    return h.hexdigest()

def read_text_from_file(path: pathlib.Path) -> str:
    ext = path.suffix.lower()
    try:
        # --- PDFs ---
        if ext == ".pdf":
            txt = []
            reader = PdfReader(str(path))
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    txt.append(t)
            return "\n\n".join(txt)

        # --- Plain / Markdown ---
        if ext in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")

        # --- Word (DOCX) ---
        if ext == ".docx" and DOCX_AVAILABLE:
            d = docx.Document(str(path))
            parts: List[str] = [p.text for p in d.paragraphs if p.text.strip()]
            for t in d.tables:
                for row in t.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n\n".join(parts)

        # --- Images (OCR) ---
        if ext in IMG_EXTS and OCR_AVAILABLE:
            img = Image.open(path)
            return pytesseract.image_to_string(img)

        # --- draw.io XML ---
        if ext == ".drawio":
            try:
                root = ET.parse(str(path)).getroot()
                texts: List[str] = []
                for elem in root.iter():
                    v = elem.attrib.get("value")
                    if v:
                        texts.append(v)
                    if elem.text and elem.text.strip():
                        texts.append(elem.text.strip())
                clean = [" ".join(t.replace("<br>", "\n").split()) for t in texts]
                return "\n\n".join(clean)
            except Exception as e:
                logging.warning(".drawio parse warning for %s: %s", path, e)
                return ""

        # --- Excel (.xlsx / .xls) ---
        if ext in {".xlsx", ".xls"}:
            try:
                import pandas as pd
                # Διαβάζουμε όλα τα sheets -> dict {sheet_name: DataFrame}
                sheets = pd.read_excel(
                    str(path),
                    sheet_name=None,
                    dtype=str,
                    engine="openpyxl" if ext == ".xlsx" else None,
                )
                parts: List[str] = []
                for sheet_name, df in sheets.items():
                    if df is None or df.size == 0:
                        continue
                    df = df.fillna("")
                    # Soft limits για τεράστια αρχεία
                    max_rows, max_cols = 2000, 50
                    if len(df) > max_rows:
                        df = df.iloc[:max_rows, :]
                    if df.shape[1] > max_cols:
                        df = df.iloc[:, :max_cols]
                    rows = [" | ".join(map(str, row)) for row in df.astype(str).values.tolist()]
                    sheet_block = f"[Sheet: {sheet_name}]\n" + "\n".join(rows)
                    parts.append(sheet_block)
                return "\n\n".join(parts)
            except Exception as e:
                logging.error("Failed to read Excel %s: %s", path, e)
                return ""

        return ""
    except Exception as e:
        logging.error("Failed to read %s: %s", path, e)
        return ""

def simple_chunk(text: str, max_chars: int = 1800, overlap: int = 200) -> List[str]:
    """Paragraph-aware chunking; robust to files without blank lines."""
    if not text:
        return []
    if "\n\n" in text:
        paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    else:
        step = 500
        paras = [text[i:i + step] for i in range(0, len(text), step)]
    chunks: List[str] = []
    buf: List[str] = []
    cur_len = 0
    for p in paras:
        if len(p) > max_chars:
            if buf:
                chunks.append("\n\n".join(buf))
                buf, cur_len = [], 0
            for i in range(0, len(p), max_chars):
                chunks.append(p[i:i + max_chars])
            continue
        if cur_len + len(p) + 2 <= max_chars:
            buf.append(p)
            cur_len += len(p) + 2
        else:
            if buf:
                chunks.append("\n\n".join(buf))
            buf, cur_len = [p], len(p)
    if buf:
        chunks.append("\n\n".join(buf))
    if overlap > 0 and len(chunks) > 1:
        out = []
        for i, c in enumerate(chunks):
            if i == 0:
                out.append(c)
            else:
                prev_tail = chunks[i - 1][-overlap:]
                out.append(prev_tail + "\n" + c)
        return out
    return chunks

# ---------- Store (with full text) ----------
def load_store() -> Tuple[List[Dict[str, Any]], np.ndarray]:
    meta: List[Dict[str, Any]] = []
    if NODES_PATH.exists():
        with NODES_PATH.open("r", encoding="utf-8") as f:
            for line in f:
                meta.append(json.loads(line))
    embs = np.load(EMB_PATH) if EMB_PATH.exists() else np.zeros((0, EMBED_DIM), dtype=np.float32)
    if embs.size and embs.shape[1] != EMBED_DIM:
        logging.error("Embedding dimension mismatch (have %d, need %d). Run 'ingest --rebuild'.", embs.shape[1], EMBED_DIM)
        sys.exit(2)
    return meta, embs

def save_store(meta: List[Dict[str, Any]], embs: np.ndarray) -> None:
    STORE_DIR.mkdir(parents=True, exist_ok=True)
    with NODES_PATH.open("w", encoding="utf-8") as f:
        for m in meta:
            f.write(json.dumps(m, ensure_ascii=False) + "\n")
    np.save(EMB_PATH, embs.astype(np.float32))

# ---------- Embeddings ----------
def enable_local_embeddings() -> None:
    """Enable local sentence-transformers and ignore any bad HF tokens."""
    global LOCAL_EMBEDDINGS, local_model, EMBED_DIM
    for var in ("HF_TOKEN", "HUGGINGFACEHUB_API_TOKEN", "HUGGINGFACE_HUB_TOKEN"):
        if os.environ.get(var):
            logging.warning("Ignoring HuggingFace token in %s for public models.", var)
            os.environ.pop(var, None)
    candidates = [
        "sentence-transformers/all-MiniLM-L6-v2",
        "all-MiniLM-L6-v2",
        "sentence-transformers/paraphrase-MiniLM-L6-v2",
        "paraphrase-MiniLM-L6-v2",
    ]
    cache_dir = str((pathlib.Path(".") / ".hf_cache").resolve())
    last_err: Optional[Exception] = None
    for name in candidates:
        try:
            local_model = SentenceTransformer(name, cache_folder=cache_dir)
            LOCAL_EMBEDDINGS = True
            EMBED_DIM = 384
            logging.info("Local embeddings enabled (%s, dim=384).", name)
            return
        except Exception as e:
            last_err = e
            logging.warning("Could not load '%s' (%s). Trying next…", name, e)
    raise RuntimeError(f"Failed to load a local embedding model. Last error: {last_err}")

def init_client() -> None:
    global client, LOCAL_EMBEDDINGS, local_model, EMBED_DIM
    load_dotenv()
    force_local = os.getenv("HANDOVER_FORCE_LOCAL", "0") in ("1", "true", "True")
    api_key = os.getenv("OPENAI_API_KEY", "")
    if not force_local and api_key and OpenAI is not None:
        client = OpenAI()
        LOCAL_EMBEDDINGS = False
        EMBED_DIM = 1536
        logging.info("OpenAI embeddings enabled (text-embedding-3-small). Set HANDOVER_FORCE_LOCAL=1 to force local.")
    else:
        if not LOCAL_EMBED_AVAILABLE:
            raise RuntimeError("Local embeddings unavailable. Install sentence-transformers or set OPENAI_API_KEY.")
        enable_local_embeddings()

def embed_texts(texts: List[str]) -> np.ndarray:
    if not texts:
        return np.zeros((0, EMBED_DIM), dtype=np.float32)

    if LOCAL_EMBEDDINGS:
        vecs = local_model.encode(texts, normalize_embeddings=False, convert_to_numpy=True)
        return vecs.astype(np.float32)

    # OpenAI path with retry/backoff; fallback to local if needed
    BATCH = 32
    out: List[np.ndarray] = []
    for i in range(0, len(texts), BATCH):
        batch = texts[i:i + BATCH]
        backoff = 0.8
        attempts = 0
        while True:
            try:
                resp = client.embeddings.create(model="text-embedding-3-small", input=batch)
                vecs = [np.array(d.embedding, dtype=np.float32) for d in resp.data]
                out.append(np.vstack(vecs))
                break
            except Exception as e:
                attempts += 1
                msg = str(e)
                if ("429" in msg) or ("quota" in msg) or ("insufficient_quota" in msg):
                    logging.warning("Embeddings error (%s). Backoff %.1fs (attempt %d)…", e.__class__.__name__, backoff, attempts)
                    time.sleep(min(backoff, 10.0))
                    backoff *= 1.8
                    if attempts >= 4:
                        if LOCAL_EMBED_AVAILABLE:
                            logging.warning("Switching to local embeddings fallback.")
                            enable_local_embeddings()
                            return embed_texts(texts)
                        raise
                else:
                    raise
    return np.vstack(out)

def cosine_top_k(query_vec: np.ndarray, matrix: np.ndarray, k: int = 5) -> List[int]:
    if matrix.size == 0:
        return []
    q = query_vec / (np.linalg.norm(query_vec) + 1e-9)
    M = matrix / (np.linalg.norm(matrix, axis=1, keepdims=True) + 1e-9)
    sims = (M @ q)
    k = min(k, len(sims))
    idx = np.argpartition(-sims, k - 1)[:k]
    idx = idx[np.argsort(-sims[idx])]
    return idx.tolist()

# ---------- Ingestion ----------
def ingest_folder(folder: str) -> None:
    base = pathlib.Path(folder)
    if not (base.exists() and base.is_dir()):
        print(f"Folder not found: {folder}")
        return

    files = [p for p in base.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS]
    if not files:
        print("No supported files.")
        return

    meta, embs = load_store()

    new_nodes: List[Dict[str, Any]] = []
    new_chunks: List[str] = []
    for p in tqdm(files, desc="Reading & chunking"):
        text = read_text_from_file(p)
        if not text.strip():
            continue
        chunks = simple_chunk(text)
        for j, ch in enumerate(chunks):
            node_id = sha1(f"{p.resolve()}::{j}::{len(ch)}")
            new_nodes.append({
                "id": node_id,
                "source": str(p.resolve()),
                "chunk_index": j,
                "char_count": len(ch),
                "preview": ch[:200].replace("\n", " ") + ("…" if len(ch) > 200 else ""),
                "text": ch,  # store FULL TEXT
            })
            new_chunks.append(ch)

    if not new_chunks:
        print("No text found.")
        return

    new_embs = embed_texts(new_chunks)
    all_meta = meta + new_nodes
    all_embs = new_embs if embs.size == 0 else np.vstack([embs, new_embs])
    save_store(all_meta, all_embs)
    print(f"Ingested {len(new_chunks)} chunks from {len(files)} files.")

# ---------- Q&A ----------
def sentence_split(text: str) -> List[str]:
    return re.split(r"(?<=[.!?])\s+", text.strip()) if text.strip() else []

def extractive_answer(question: str, ctx_chunks: List[Dict[str, Any]]) -> str:
    """Offline extractive: pull a few sentences from top chunks that contain query terms."""
    # Small rule for RFC1918 quick answers
    m = re.search(r"(\d{1,3}(?:\.\d{1,3}){3})", question)
    if m:
        ip = m.group(1)
        parts = ip.split(".")
        if len(parts) == 4:
            try:
                p0, p1 = int(parts[0]), int(parts[1])
                if p0 == 10 or (p0 == 172 and 16 <= p1 <= 31) or (p0 == 192 and p1 == 168):
                    return f"Η IP {ip} ανήκει σε **ιδιωτικό IPv4 χώρο (RFC1918)**. Συνήθως χρησιμοποιείται σε LAN (gateway/management, DHCP, συσκευές)."
            except Exception:
                pass

    q = question.lower()
    terms = [w for w in re.findall(r"\w+", q) if len(w) >= 4]
    scored: List[str] = []
    for m in ctx_chunks:
        sents = sentence_split(m.get("text", ""))
        for s in sents:
            score = sum(1 for t in terms if t in s.lower())
            if score > 0:
                scored.append(s)
    # Dedup & top few
    seen = set()
    dedup: List[str] = []
    for s in scored:
        k = s.strip()
        if k not in seen:
            seen.add(k)
            dedup.append(k)
        if len(dedup) >= 6:
            break
    if dedup:
        bullets = [f"• {x}" for x in dedup]
        return "\n".join(bullets)
    lines = [f"• Δες το αρχείο: {os.path.basename(m['source'])} (chunk {m['chunk_index']})" for m in ctx_chunks]
    return "\n".join(lines) if lines else "Δεν βρέθηκε σχετικό απόσπασμα. Δες τις παραπομπές."

def answer_query(question: str, k: int = 5) -> None:
    meta, embs = load_store()
    if not meta:
        print("Empty store, run ingest first.")
        return

    q_vec = embed_texts([question])[0]
    idxs = cosine_top_k(q_vec, embs, k)
    picked = [meta[idx] for idx in idxs]

    ctx_lines: List[str] = []
    for i, m in enumerate(picked, start=1):
        full = m.get("text", "") or ""
        # δώσε στο μοντέλο αρκετό περιεχόμενο, χωρίς να ξεφύγεις
        SNIPPET_MAX = 1500
        snippet = full[:SNIPPET_MAX] + ("…" if len(full) > SNIPPET_MAX else "")
        ctx_lines.append(f"[{i}] Source: {os.path.basename(m['source'])}\n{snippet}")
    ctx = "\n".join(ctx_lines)


    use_openai_chat = os.getenv("HANDOVER_FORCE_LOCAL", "0") not in ("1", "true", "True") and client is not None

    if use_openai_chat:
        try:
            prompt = f"""Answer the question strictly from the context below.
If unsure, say you don't have enough information and point to the most relevant sources.
Prefer bullet points and include citations like [#].

Context:
{ctx}

Question: {question}"""
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
            )
            content = resp.choices[0].message.content
        except Exception as e:
            msg = str(e)
            if ("429" in msg) or ("quota" in msg) or ("insufficient_quota" in msg):
                logging.warning("Chat quota/429 – switching to offline extractive answer.")
                content = extractive_answer(question, picked)
            else:
                raise
    else:
        content = extractive_answer(question, picked)

    print("=== Answer ===")
    print(content)
    print("--- Citations ---")
    for i, m in enumerate(picked, start=1):
        print(f"[{i}] {m['source']} (chunk {m['chunk_index']})")

# ---------- CLI ----------
def main() -> None:
    ap = argparse.ArgumentParser()
    sub = ap.add_subparsers(dest="cmd")

    p1 = sub.add_parser("ingest")
    p1.add_argument("folder")
    p1.add_argument("--rebuild", action="store_true", help="Rebuild the store from scratch")

    p2 = sub.add_parser("ask")
    p2.add_argument("question")

    args = ap.parse_args()
    init_client()

    if args.cmd == "ingest":
        if getattr(args, "rebuild", False) and STORE_DIR.exists():
            shutil.rmtree(STORE_DIR)
        ingest_folder(args.folder)
    elif args.cmd == "ask":
        answer_query(args.question)
    else:
        ap.print_help()

if __name__ == "__main__":
    main()
